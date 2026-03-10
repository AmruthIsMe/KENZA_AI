#!/usr/bin/env python3
"""
Generate IEEE Conference Paper - Sections III & IV
Corrected and verified against actual Kenza source code.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1, numbering=None):
    """Add a formatted heading"""
    p = doc.add_paragraph()
    p.space_before = Pt(14 if level == 1 else 10)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 3:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.italic = True
    return p


def add_paragraph(doc, text, bold_prefix=None, indent=False):
    """Add justified paragraph with optional bold prefix"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    for run in p.runs:
        run.font.size = Pt(10)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p


def add_bullet(doc, text, indent=False):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    if indent:
        p.paragraph_format.left_indent = Inches(0.5)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = 1  # center

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'D9E2F3')
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)

    for r, row in enumerate(rows):
        row_cells = table.rows[r + 1].cells
        for c, val in enumerate(row):
            row_cells[c].text = str(val)
            for para in row_cells[c].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")
    return table


def generate_ieee_paper():
    doc = Document()

    # Configure default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    # =========================================================================
    # III. SYSTEM ARCHITECTURE
    # =========================================================================

    add_heading(doc, "III. SYSTEM ARCHITECTURE")

    add_paragraph(doc,
        "The Kenza telepresence robot follows a modular embedded architecture "
        "integrating computation, user interaction, artificial intelligence services, "
        "and motor actuation within a unified framework. The system is designed around "
        "a central processing unit that manages perception, communication, and control "
        "while interfacing with cloud-based AI services and a browser-based user application. "
        "The architecture is divided into four major subsystems: Hardware Platform, "
        "Web Application Interface, AI Agent Framework, and Motor Control System."
    )

    # --- A. Hardware Component ---
    add_heading(doc, "A. Hardware Platform", level=2)

    add_paragraph(doc,
        "The Raspberry Pi 5 (8GB RAM) functions as the sole computational node of the system. "
        "It manages camera acquisition via the Raspberry Pi Camera Module 3, Canvas-based animated "
        "eye rendering on an HDMI-connected display, AI request processing through cloud APIs and "
        "local model inference, GPIO-based motor control through the L298N dual H-bridge driver, "
        "and hosting of all server processes including the WebSocket server (port 8765) and HTTP "
        "motor server (port 8080). By centralizing all processing tasks within a single embedded "
        "platform, system complexity is minimized while maintaining high functional capability."
    )

    add_paragraph(doc,
        "The Raspberry Pi Camera Module 3 captures real-time video at 640\u00d7480 resolution "
        "(configurable up to 1080p) for telepresence streaming via WebRTC and contextual AI "
        "understanding through Gemini's multimodal vision API. The HDMI-connected display provides "
        "immediate visual feedback, rendering animated eye expressions and operational mode "
        "indicators through a Chromium kiosk-served web interface (robot_display.html), which "
        "contributes to a socially expressive robotic presence. The L298N dual H-bridge module "
        "interfaces directly with four Raspberry Pi GPIO pins (GPIO17, GPIO27, GPIO22, GPIO23) "
        "using the gpiozero library for Pi 5 compatibility, controlling the DC geared motors "
        "responsible for differential-drive locomotion. Three 18650 lithium-ion batteries "
        "connected in parallel ensure stable and sustained power delivery, particularly during "
        "motor actuation, where higher torque demand is observed."
    )

    add_paragraph(doc,
        "Status indication is provided by two GPIO-controlled LEDs: a Green LED (GPIO24/BCM Pin 24) "
        "indicating the listening/ready state, and a Red LED (GPIO25/BCM Pin 25) indicating the "
        "thinking/processing state. These LEDs are managed by the LEDController class within the "
        "AI subsystem."
    )

    add_paragraph(doc,
        "From a human-centered perspective, the hardware platform is not merely a collection of "
        "components; it forms the physical embodiment of the AI companion, enabling perception "
        "(Camera Module 3), expression (HDMI display + USB speaker), cognition (Pi 5 + Cloud AI "
        "APIs + Local LLMs), and action (motor system)."
    )

    # Table: Hardware
    add_table(doc,
        ["Component", "Specification", "Purpose"],
        [
            ["Compute Unit", "Raspberry Pi 5 (8GB LPDDR4X)", "Core processing, AI, Control"],
            ["Motor Driver", "L298N Dual H-Bridge", "DC Motor Direction Control"],
            ["Camera", "Raspberry Pi Camera Module 3 (CSI-2)", "1080p Streaming, Vision AI"],
            ["Audio Input", "USB Microphone", "Speech Recognition, Streaming"],
            ["Audio Output", "USB Speaker / 3.5mm Jack", "TTS, Telepresence Audio"],
            ["Display", "HDMI Display (Chromium Kiosk)", "Eye Animations, Status Modes"],
            ["Status LEDs", "Green (GPIO24), Red (GPIO25)", "Listening / Thinking Indicators"],
            ["Power", "3x 18650 Li-ion (Parallel)", "System + Motor Power"],
            ["Connectivity", "Wi-Fi + WebSocket + WebRTC", "Comms, Telemetry, Streaming"],
        ]
    )

    # --- B. Web Application Interface ---
    add_heading(doc, "B. Web Application Interface", level=2)

    add_paragraph(doc,
        "The Kenza system adopts a dual-screen interaction model, operating simultaneously "
        "across the onboard HDMI display and a browser-based remote user interface. This "
        "architecture enhances telepresence by creating synchronized interaction between the "
        "robot body and the remote user device. This dual-interface design strengthens user "
        "immersion by aligning physical robot expression with remote command control."
    )

    add_heading(doc, "1) Robot Display (Onboard Interface):", level=3)

    add_paragraph(doc,
        "The onboard interface is rendered via robot_display.html, served in Chromium kiosk mode "
        "on the Pi's HDMI display. It represents the robot's \"face\" and operational dashboard. "
        "It operates in three primary modes:"
    )

    add_bullet(doc,
        "Action Mode: Displays illuminating animated eyes synchronized with AI conversation. "
        "Users can initiate AI chat sessions or contact-based telepresence calling. The eye "
        "animations change dynamically to reflect listening, thinking, speaking, happy, sad, "
        "excited, and confused states, reinforcing emotional engagement. The EmotionEyeBridge "
        "class sends real-time emotion state signals via WebSocket at ~12 Hz polling."
    )
    add_bullet(doc,
        "Explore Mode: Activates sentry-based patrol behavior. The system monitors environmental "
        "input using YOLO11n object detection (offline) or MediaPipe face detection and generates "
        "WebSocket-based alerts when unfamiliar human presence is detected. The display visually "
        "communicates active surveillance status."
    )
    add_bullet(doc,
        "Settings Mode: Provides onboard configuration including QR-based Wi-Fi pairing, "
        "connection status display, and volume adjustment sliders, enabling localized user control."
    )

    add_heading(doc, "2) User Web Application (Browser-Based Interface):", level=3)

    add_paragraph(doc,
        "The web application (kenza_app_v2.html), developed as a standalone HTML/CSS/JavaScript "
        "Single-Page Application, serves as the primary remote control interface. Communication is "
        "handled via two channels: a WebSocket connection (port 8765, asyncio websockets library) for "
        "bidirectional low-latency control, telemetry, and AI messaging at approximately 20 Hz; and "
        "WebRTC (via aiortc on the Pi) for 720p real-time video streaming with VP8/VP9 encoding. "
        "A secondary HTTP fallback path (port 8080) provides motor command capability. "
        "Key features include:"
    )

    add_bullet(doc, "Manual virtual joystick for precise directional movement")
    add_bullet(doc, "WASD keyboard controls and touch-based directional input")
    add_bullet(doc, "Eye customization (color and expression personalization)")
    add_bullet(doc, "Telepresence video call interface with WebRTC signaling relay")
    add_bullet(doc, "Live telemetry dashboard (battery, Wi-Fi RSSI, CPU temperature)")
    add_bullet(doc, "AI chat text input with real-time response display")
    add_bullet(doc, "Voice preset selection (5 TTS personalities)")
    add_bullet(doc, "Sentry mode controls with alert notification UI")

    # --- C. AI Companion Agents ---
    add_heading(doc, "C. AI Companion Agents", level=2)

    add_paragraph(doc,
        "The conversational AI subsystem employs a multi-model routing architecture with "
        "automatic cloud/offline failover, implemented in the ConversationEngine class "
        "(kenza_conversation.py). The system chains four LLM providers in priority order:"
    )

    add_table(doc,
        ["Priority", "Provider", "Model", "Use Case"],
        [
            ["1 (Primary)", "Groq", "llama-3.3-70b-versatile", "Fast conversational inference (~800-2000 ms)"],
            ["2 (Vision)", "Gemini", "gemini-2.0-flash", "Multimodal vision queries, secondary text fallback"],
            ["3 (Offline)", "Ollama", "gemma3:270m (configurable)", "Local server, zero-internet operation"],
            ["4 (Offline)", "LlamaChat", "llama-3.2-3b-instruct.Q4_K_M", "GGUF file fallback via llama-cpp-python"],
        ]
    )

    add_paragraph(doc,
        "Both cloud providers (Groq and Gemini) execute API calls from the Raspberry Pi 5. "
        "Audio input is captured via PyAudio (16 kHz, mono) and processed by a Voice Activity "
        "Detector (VAD) using RMS energy thresholding with sustained-frame validation to handle "
        "speech interruptions during TTS playback. Speech is transcribed using Google's Speech "
        "Recognition API (online) with automatic fallback to faster-whisper (base.en model, "
        "offline). A ConnectivityMonitor class performs periodic socket pings (8.8.8.8:53) every "
        "30 seconds to determine cloud/offline routing automatically."
    )

    add_paragraph(doc,
        "Generated responses are processed through the KenzaPersonality module, which enforces "
        "emotion-tagged replies (e.g., [happy], [thinking], [sad]). The EmotionEngine extracts "
        "these tags and maps them to corresponding eye animation states on the HDMI display. "
        "Text-to-speech conversion uses Microsoft Edge-TTS (online, neural voices) with automatic "
        "fallback to pyttsx3/espeak-ng (offline). Emotional prosody is applied via Edge-TTS "
        "rate and pitch parameters, creating a human-like \"AI mindset\" with companion-like "
        "interaction."
    )

    add_paragraph(doc, "", bold_prefix="AI Workflow: ")
    add_paragraph(doc,
        "Voice Input \u2192 VAD Filter \u2192 Google STT / Whisper \u2192 "
        "Command Parser \u2192 Groq/Gemini/Ollama Routing \u2192 Emotion-Tagged Response \u2192 "
        "Edge-TTS / espeak \u2192 Audio Playback + Eye Animation Sync"
    )

    # --- D. Motor Control System ---
    add_heading(doc, "D. Motor Control System", level=2)

    add_paragraph(doc,
        "The locomotion mechanism employs a differential-drive configuration with two DC geared "
        "motors and a caster wheel, enabling point-turn steering. The motors are controlled via "
        "the GPIOMotorController class (RPI_kenza_main.py) using gpiozero OutputDevice instances "
        "connected to the L298N H-bridge module. The control operates in binary on/off mode "
        "(direction-only), with directional commands dispatched via both WebSocket and HTTP interfaces."
    )

    add_heading(doc, "GPIO Pin Mapping:", level=3)

    add_table(doc,
        ["Signal", "GPIO Pin", "Physical Pin", "L298N Connection", "Function"],
        [
            ["IN1", "GPIO17", "Pin 11", "L298N IN1", "Motor A Forward"],
            ["IN2", "GPIO27", "Pin 13", "L298N IN2", "Motor A Reverse"],
            ["IN3", "GPIO22", "Pin 15", "L298N IN3", "Motor B Forward"],
            ["IN4", "GPIO23", "Pin 16", "L298N IN4", "Motor B Reverse"],
        ]
    )

    add_heading(doc, "Directional Logic:", level=3)

    add_table(doc,
        ["Command", "IN1", "IN2", "IN3", "IN4", "Result"],
        [
            ["Forward (F)", "ON", "OFF", "ON", "OFF", "Both motors forward"],
            ["Backward (B)", "OFF", "ON", "OFF", "ON", "Both motors reverse"],
            ["Left (L)", "OFF", "ON", "ON", "OFF", "Pivot left (differential)"],
            ["Right (R)", "ON", "OFF", "OFF", "ON", "Pivot right (differential)"],
            ["Stop (S)", "OFF", "OFF", "OFF", "OFF", "All motors halted"],
        ]
    )

    add_paragraph(doc,
        "Three parallel-connected 18650 lithium-ion batteries (3.7V each) provide consistent "
        "voltage supply for high-torque motor operation. A voltage regulator ensures stable "
        "5V and 3.3V rails for the Raspberry Pi and peripheral components. This configuration "
        "ensures extended operational duration and stable mobility performance."
    )

    add_paragraph(doc, "Safety mechanisms include:")
    add_bullet(doc, "Emergency stop command accessible from both WebSocket and HTTP interfaces")
    add_bullet(doc, "Automatic GPIO cleanup on process termination (gpiozero resource management)")
    add_bullet(doc, "Motor watchdog: all motors halted when client WebSocket disconnects")

    add_paragraph(doc,
        "The motor system translates digital user intention (joystick commands or voice commands "
        "via the CommandParser) into physical movement, completing the perception-decision-action "
        "loop that defines embodied telepresence robotics."
    )

    # =========================================================================
    # IV. METHODOLOGY
    # =========================================================================

    add_heading(doc, "IV. METHODOLOGY")

    add_paragraph(doc,
        "This section presents the systematic implementation procedure of the Raspberry Pi 5-based "
        "telepresence companion robot. The development process followed an iterative and "
        "milestone-driven methodology, progressing through platform configuration, web interface "
        "development, AI integration, motor control implementation, and validation testing. "
        "Each subsystem was independently verified before full system integration to ensure "
        "modular reliability and reproducibility."
    )

    # --- A. Platform Configuration ---
    add_heading(doc, "A. Platform Configuration and Setup", level=2)

    add_paragraph(doc,
        "The system was initialized using the Raspberry Pi 5 (8GB RAM) as the primary and sole "
        "computational unit. Raspberry Pi OS Bookworm (64-bit) was installed and optimized for "
        "real-time video streaming, GPIO motor control, and API-based AI service integration. "
        "The system boots via a single unified entry point (RPI_kenza_main.py) which automatically "
        "starts all services:"
    )

    add_bullet(doc, "WebSocket Server (port 8765) \u2014 asyncio websockets library for bidirectional control and telemetry")
    add_bullet(doc, "HTTP Motor Server (port 8080) \u2014 lightweight BaseHTTPRequestHandler for motor command fallback")
    add_bullet(doc, "WebRTC Streaming \u2014 aiortc library for VP8/VP9 video + Opus audio")
    add_bullet(doc, "Chromium Kiosk Display \u2014 robot_display.html rendered fullscreen on HDMI output")
    add_bullet(doc, "Conversation/Voice Engine \u2014 ConversationEngine with wake word detection")
    add_bullet(doc, "Telemetry Broadcaster \u2014 periodic system metrics (CPU temp, Wi-Fi RSSI, battery) via WebSocket")

    add_heading(doc, "1) GPIO Motor Control Mapping:", level=3)

    add_paragraph(doc,
        "Motor driver interfacing was configured using the gpiozero library (OutputDevice class) "
        "for Raspberry Pi 5 compatibility. The following GPIO mapping was established:"
    )

    add_bullet(doc, "GPIO17 (Pin 11) \u2192 L298N IN1 \u2014 Motor A direction control")
    add_bullet(doc, "GPIO27 (Pin 13) \u2192 L298N IN2 \u2014 Motor A direction control")
    add_bullet(doc, "GPIO22 (Pin 15) \u2192 L298N IN3 \u2014 Motor B direction control")
    add_bullet(doc, "GPIO23 (Pin 16) \u2192 L298N IN4 \u2014 Motor B direction control")

    add_paragraph(doc,
        "This mapping allows the Raspberry Pi to directly regulate motor direction through "
        "binary GPIO signals. Differential steering is achieved by selectively activating "
        "opposing motor pairs (e.g., left motor forward + right motor reverse for pivot turning)."
    )

    add_heading(doc, "2) Power Configuration:", level=3)

    add_paragraph(doc,
        "Three 18650 lithium-ion batteries (3.7V each) were connected in parallel to provide "
        "high-capacity and stable current output. A buck converter was employed to regulate "
        "voltage levels to 5V and 3.3V rails, ensuring stable operation of both the Raspberry Pi "
        "and the L298N motor driver. This configuration was chosen to maintain consistent "
        "performance during simultaneous AI processing and motor actuation, preventing voltage "
        "drops that could interrupt system operation."
    )

    # --- B. AI Companion Agent Integration ---
    add_heading(doc, "B. AI Companion Agent Integration", level=2)

    add_paragraph(doc,
        "The conversational subsystem implements a four-tier LLM pipeline with automatic "
        "failover, structured as:"
    )

    add_paragraph(doc,
        "Voice Input \u2192 VAD (RMS Energy) \u2192 Google STT / faster-whisper \u2192 "
        "Command Parser \u2192 Query Router \u2192 Groq / Gemini / Ollama / Llama \u2192 "
        "Emotion Extraction \u2192 Edge-TTS / pyttsx3 \u2192 Display + Eye Animation"
    )

    add_paragraph(doc,
        "Speech input is first transcribed into text. When internet is available, Google's "
        "Speech Recognition API provides high-accuracy cloud STT. When offline, faster-whisper "
        "(base.en model via CTranslate2) provides local transcription. A ConnectivityMonitor "
        "class (socket ping to 8.8.8.8:53 every 30s) transparently switches between modes."
    )

    add_paragraph(doc,
        "The ConversationEngine's process_input() method implements the following routing logic:"
    )

    add_bullet(doc, "Step 1 \u2014 Voice commands are intercepted by CommandParser for direct actions "
                    "(eye color changes, mode switches, voice preset selection, follow/explore commands).")
    add_bullet(doc, "Step 2 \u2014 Vision queries (triggered by phrases like \"what do you see\") are routed to "
                    "Gemini 2.0 Flash multimodal API (online) or YOLO11n object detection (offline).")
    add_bullet(doc, "Step 3 \u2014 General conversation is sent to Groq (llama-3.3-70b-versatile) for "
                    "ultra-fast inference. If Groq fails, the query falls through to Gemini 2.0 Flash.")
    add_bullet(doc, "Step 4 \u2014 If both cloud providers are unavailable, the system falls back to "
                    "Ollama (gemma3:270m by default, configurable) or llama-cpp-python GGUF model.")

    add_paragraph(doc,
        "The generated response is parsed for emotion tags ([happy], [sad], [excited], [neutral], "
        "[confused], [thinking]) by the _parse_and_strip_emotion() method. These emotion states "
        "are broadcasted in real-time via the EmotionEyeBridge to the HDMI display, where "
        "Canvas-based eye animations synchronize with the verbal response. The response text is "
        "converted to speech using Microsoft Edge-TTS (neural voices, with emotion-driven prosody "
        "adjustment for rate and pitch) or pyttsx3/espeak-ng when offline."
    )

    add_heading(doc, "1) Sentry Mode Logic:", level=3)

    add_paragraph(doc,
        "The sentry module operates through the following pipeline:"
    )

    add_paragraph(doc,
        "Autonomous patrol \u2192 Camera frame acquisition \u2192 "
        "Face/Human detection (MediaPipe + YOLO11n) \u2192 "
        "Stranger identification (face_recognition/dlib) \u2192 "
        "WebSocket alert broadcast to connected user application"
    )

    add_paragraph(doc,
        "When a human presence is detected during patrol mode, the system triggers real-time "
        "WebSocket notification messages to all connected controller clients. The web application "
        "displays visual and audio alerts, and can optionally initiate video recording. This mode "
        "transforms the robot from a passive telepresence device into an active monitoring "
        "assistant. The motor subsystem translates autonomous navigation decisions into physical "
        "movement, completing the perception-decision-action cycle essential to telepresence robotics."
    )

    # --- C. Integration Testing ---
    add_heading(doc, "C. Integration Testing and Validation", level=2)

    add_table(doc,
        ["Phase", "Component", "Test Case", "Success Criteria"],
        [
            ["1", "Motor Control", "GPIO signal verification for all 5 directions", "Correct ON/OFF per command"],
            ["2", "WebSocket Server", "Bidirectional message relay + state sync", "< 50 ms round-trip latency"],
            ["3", "WebRTC Streaming", "VP8 video + Opus audio via aiortc", "720p @ 30 FPS, < 300 ms latency"],
            ["4", "STT Pipeline", "Google STT (online) + Whisper (offline) switching", "\u2265 90% word accuracy"],
            ["5", "AI Routing", "Groq \u2192 Gemini \u2192 Ollama \u2192 Llama failover chain", "Response in all connectivity states"],
            ["6", "TTS Pipeline", "Edge-TTS (online) + pyttsx3 (offline) + emotion prosody", "Audible, emotion-matched output"],
            ["7", "Eye Animations", "Emotion tag \u2192 Display sync via EmotionEyeBridge", "\u2264 100 ms tag-to-animation delay"],
            ["8", "Sentry Mode", "Human detection \u2192 WebSocket alert broadcast", "Alert within 2 seconds of detection"],
            ["9", "Full Integration", "End-to-end user workflow test", "All subsystems operational concurrently"],
        ]
    )

    add_heading(doc, "1) Demonstration Workflow:", level=3)

    add_paragraph(doc,
        "User web app connection \u2192 Joystick motor control \u2192 Live video feedback via WebRTC \u2192 "
        "AI chat interaction (voice or text) \u2192 Eye animation response \u2192 Sentry mode activation \u2192 "
        "Human detection alert notification"
    )

    add_heading(doc, "2) Performance Targets:", level=3)

    add_table(doc,
        ["Metric", "Target", "Measured"],
        [
            ["Joystick-to-motor latency", "< 200 ms", "~ 50-150 ms (WebSocket)"],
            ["AI response coherence", "\u2265 95%", "Groq: consistent at scale"],
            ["AI response latency (Groq)", "< 2000 ms", "800-2000 ms typical"],
            ["AI response latency (Local)", "< 5000 ms", "2000-5000 ms (Llama 3.2)"],
            ["WebRTC video latency", "< 300 ms", "100-300 ms (VP8)"],
            ["Wake word detection", "< 200 ms", "< 200 ms (Google STT)"],
            ["Continuous operation", "\u2265 2 hours", "2+ hours (3x 18650)"],
            ["Face detection (per frame)", "< 30 ms", "15-30 ms (MediaPipe)"],
        ]
    )

    add_paragraph(doc,
        "The iterative validation strategy ensured subsystem reliability prior to integration, "
        "resulting in a functional prototype capable of seamless telepresence interaction and "
        "AI-assisted companionship."
    )

    # Save
    filename = "Kenza_IEEE_Paper_Sections.docx"
    doc.save(filename)
    print(f"Successfully generated {filename}")


if __name__ == "__main__":
    generate_ieee_paper()
