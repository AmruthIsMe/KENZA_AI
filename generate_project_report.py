import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(46, 116, 181)  # Blue
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(46, 116, 181)
    else:
        run.font.size = Pt(12)
    return p

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'D9E2F3')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Rows
    for r, row in enumerate(rows):
        row_cells = table.rows[r + 1].cells
        for c, val in enumerate(row):
            row_cells[c].text = str(val)
            
    return table

def generate_report():
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Kenza Project: Technical Specification and Feature Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 50, 100)
    
    doc.add_paragraph("\n")
    
    # 1. Project Overview
    add_heading(doc, "1. Project Overview")
    add_paragraph(doc, 
        "Kenza is an advanced AI-powered telepresence robot designed for remote interaction, autonomous assistance, and social engagement. "
        "Built on a Raspberry Pi 5 platform, it integrates multimodal AI (Vision, Voice, LLM), real-time WebRTC streaming, and autonomous navigation capabilities into a cohesive, low-cost robotic agent."
    )
    
    # 2. Key Features
    add_heading(doc, "2. Key Features")
    
    features = [
        "Two-way AV communication: Full-duplex audio/video streaming with Acoustic Echo Cancellation (AEC).",
        "Real-time video streaming: Low-latency WebRTC (VP8/VP9) via PiCamera2.",
        "Obstacle Avoidance: Zone-based collision detection using MobileNet-SSD or contour analysis.",
        "Edge Processing: All core logic (Motor control, STT, TTS, Vision) runs locally on Raspberry Pi 5.",
        "Local IP Based: Zero-configuration local network discovery and control.",
        "Expressive Animated Eyes: TFT display rendering emotional states (Blinking, Thinking, Happy, etc.) synced with voice.",
        "Gesture Control: MediaPipe-based hand tracking for driving and UI interaction (11 supported gestures).",
        "Human Following: Autonomous person tracking using Pose estimation and Histograms for re-identification.",
        "Wake Word Activation: \"Kenza\" wake word detection using Google Speech Recognition.",
        "Conversation Engine: Tri-model pipeline routing queries to Cloud (Groq/Llama-3.3, Gemini 2.0) or Local (Llama 3.2 3B) models.",
        "Autonomous Exploration: Random-walk exploration with obstacle avoidance behavior.",
        "Multiple Avatar Selection: 5 distinct TTS voice presets (Kenza, Glitch, Kawaii, Titan, Jarvis).",
        "Overnight Surveillance: 'Sentry Mode' utilizing object detection for monitoring (Human detection triggers).",
        "Web App Control: Responsive Mobile/Desktop web UI with Joystick, WASD keys, and touch support.",
        "Display Visuals: Front-facing screen for telepresence feedback and emotional avatars.",
        "Movement: Differential drive (2 Wheels + Caster) or Tank-style movement logic."
    ]
    
    for f in features:
        add_bullet(doc, f)

    # 3. Technical Specifications
    add_heading(doc, "3. Technical Specifications")
    
    # 3.1 Hardware
    add_heading(doc, "3.1 Hardware Components", level=2)
    hw_data = [
        ["Component", "Specification", "Purpose"],
        ["Compute Unit", "Raspberry Pi 5 (8GB RAM)", "Core processing, AI, Control"],
        ["Motor Driver", "L298N", "Dual H-Bridge DC Motor Control"],
        ["Camera", "Pi Camera Module 3 / PiCamera2", "1080p Streaming, Vision AI"],
        ["Audio Input", "USB Microphone", "Speech Recognition, Streaming"],
        ["Audio Output", "USB Speaker / 3.5mm Jack", "TTS, Streaming Audio"],
        ["Display", "TFT/LCD Screen", "Eye Animations, Status"],
        ["Chassis", "Custom 3D Printed + Metal Base", "Structure, Enclosure"],
        ["Connectivity", "Wi-Fi + WebSocket", "Comms, Telemetry, WebRTC"]
    ]
    add_table(doc, hw_data[0], hw_data[1:])
    
    doc.add_paragraph("\n")
    
    # 3.2 Software Stack
    add_heading(doc, "3.2 Software & AI Stack", level=2)
    sw_data = [
        ["Module", "Technology / Library", "Version / Details"],
        ["OS", "Raspberry Pi OS (Bookworm)", "64-bit"],
        ["Language", "Python", "3.11+"],
        ["Speech-to-Text (SST)", "SpeechRecognition (Google API)", ">=3.10.0"],
        ["Text-to-Speech (TTS)", "Edge-TTS (Microsoft Azure Neural)", ">=6.1.0"],
        ["LLM (Cloud)", "Groq (Llama-3.3-70b), Gemini 2.0 Flash", "API-based"],
        ["LLM (Local)", "Llama 3.2 3B (Quantized GGUF)", "llama-cpp-python >=0.2.0"],
        ["Vision / Face", "MediaPipe, Face_Recognition (dlib)", "mediapipe>=0.10.0"],
        ["Object Detection", "MobileNet-SSD (Caffe)", "OpenCV DNN"],
        ["Streaming", "WebRTC (aiortc)", "aiortc>=1.6.0"],
        ["Web Server", "Flask + Websockets", "Asyncio implementation"]
    ]
    add_table(doc, sw_data[0], sw_data[1:])

    # 4. Detailed Feature Working
    add_heading(doc, "4. Detailed Feature Working")
    
    # 4.1 Conversation
    add_heading(doc, "4.1 Conversation Engine", level=2)
    add_paragraph(doc, 
        "The conversation system uses a sophisticated 'Tri-Model' routing architecture. "
        "Audio is captured via PyAudio and processed by a Voice Activity Detector (RMS-based) to handle interruptions. "
        "Speech is transcribed using Google's Speech Recognition API. "
        "The text is sent to a Smart Router which decides the backend: "
        "1) Groq (Llama 3.3) for fast, general conversation. "
        "2) Gemini 2.0 Flash for vision-related queries (e.g., 'What do you see?'). "
        "3) Local Llama 3.2 if internet is unavailable. "
        "Responses are converted to audio using Microsoft Edge-TTS with selectable voice personalities."
    )
    
    # 4.2 Vision & Autonomy
    add_heading(doc, "4.2 Vision & Autonomy", level=2)
    add_paragraph(doc,
        "Vision processing runs on the PiCamera2 feed. "
        "Face Detection (MediaPipe) identifies users and tracks their position for the 'Human Following' mode, "
        "utilizing a PID-like controller to drive the motors and keep the subject centered. "
        "Face Recognition (dlib) identifies specific users (e.g. 'Owner') to customize interactions. "
        "Autonomous movement uses MobileNet-SSD to detect obstacles (chairs, people, tables) and a "
        "zone-based collision avoidance algorithm to navigate safely."
    )
    
    # 4.3 Gesture Control
    add_heading(doc, "4.3 Gesture Control", level=2)
    add_paragraph(doc,
        "Hand gestures are tracked in real-time (30 FPS) using MediaPipe Hands. "
        "Recognized gestures are mapped to robot actions: "
        "'Pinch' for clicking/selecting, 'Fist' for dragging, and 'Pointing' for directional movement. "
        "This allows for 'Air-Touch' control of the UI and robot navigation without physical contact."
    )

    # 4.4 Telepresence & Streaming
    add_heading(doc, "4.4 Telepresence & Web App", level=2)
    add_paragraph(doc,
        "The Web Application serves as the primary control interface. "
        "It establishes a WebSocket connection for low-latency control (motor commands, settings) and telemetry (battery, temp). "
        "Video and Audio are streamed via WebRTC (using aiortc on the Pi). "
        "An Acoustic Echo Cancellation (AEC) filter prevents audio feedback loops, enabling clear two-way communication."
    )

    filename = "Kenza_Project_Report.docx"
    doc.save(filename)
    print(f"Successfully generated {filename}")

if __name__ == "__main__":
    generate_report()
